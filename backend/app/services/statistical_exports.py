from __future__ import annotations

import io
import json
import re
from html import escape, unescape
from typing import Any, Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer


def _numeric(value: Any) -> float | None:
    try:
        return float(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _grouped_means(
    rows: list[dict[str, Any]], x_key: str, y_key: str
) -> tuple[list[str], list[float]]:
    groups: dict[str, list[float]] = {}
    for row in rows:
        numeric = _numeric(row.get(y_key))
        if numeric is None:
            continue
        groups.setdefault(str(row.get(x_key, "Sem grupo")), []).append(numeric)
    labels = list(groups)
    means = [sum(groups[label]) / len(groups[label]) for label in labels]
    return labels, means


def render_chart_bytes(
    chart_type: str,
    rows: list[dict[str, Any]],
    configuration: dict[str, Any],
    title: str,
    description: str,
    output_format: str = "png",
) -> bytes:
    x_key = str(configuration.get("x_key") or "group")
    y_key = str(configuration.get("y_key") or "score")
    group_key = str(configuration.get("group_key") or "group")
    figure, axis = plt.subplots(figsize=(9, 5.2), constrained_layout=True)

    if chart_type == "paired":
        before_key = str(configuration.get("before_key") or configuration.get("x_value_key") or "pre")
        after_key = str(configuration.get("after_key") or y_key or "post")
        pairs = [
            (_numeric(row.get(before_key)), _numeric(row.get(after_key))) for row in rows
        ]
        pairs = [(before, after) for before, after in pairs if before is not None and after is not None]
        for before, after in pairs:
            axis.plot([0, 1], [before, after], marker="o", alpha=0.55)
        if pairs:
            before_mean = sum(value[0] for value in pairs) / len(pairs)
            after_mean = sum(value[1] for value in pairs) / len(pairs)
            axis.plot([0, 1], [before_mean, after_mean], marker="o", linewidth=3, label="Média")
            axis.legend()
        axis.set_xticks([0, 1], [before_key, after_key])
        axis.set_ylabel("Resultado")
    elif chart_type == "scatter":
        points = [
            (_numeric(row.get(x_key)), _numeric(row.get(y_key))) for row in rows
        ]
        points = [(x_value, y_value) for x_value, y_value in points if x_value is not None and y_value is not None]
        axis.scatter([point[0] for point in points], [point[1] for point in points])
        axis.set_xlabel(x_key)
        axis.set_ylabel(y_key)
    elif chart_type == "histogram":
        values = [_numeric(row.get(y_key)) for row in rows]
        axis.hist([value for value in values if value is not None], bins="auto")
        axis.set_xlabel(y_key)
        axis.set_ylabel("Frequência")
    elif chart_type == "boxplot":
        groups: dict[str, list[float]] = {}
        for row in rows:
            value = _numeric(row.get(y_key))
            if value is not None:
                groups.setdefault(str(row.get(group_key, "Grupo")), []).append(value)
        if groups:
            axis.boxplot(list(groups.values()), tick_labels=list(groups))
        axis.set_ylabel(y_key)
    elif chart_type == "line":
        labels, means = _grouped_means(rows, x_key, y_key)
        axis.plot(labels, means, marker="o")
        axis.set_ylabel(f"Média de {y_key}")
        axis.tick_params(axis="x", rotation=30)
    else:
        labels, means = _grouped_means(rows, x_key, y_key)
        axis.bar(labels, means)
        axis.set_ylabel(f"Média de {y_key}")
        axis.tick_params(axis="x", rotation=30)

    axis.set_title(title)
    if description:
        figure.text(0.5, 0.01, description, ha="center", fontsize=9)
    axis.grid(axis="y", alpha=0.2)
    output = io.BytesIO()
    figure.savefig(output, format=output_format, dpi=200 if output_format == "png" else None)
    plt.close(figure)
    return output.getvalue()


def _plain_text(html_content: str) -> str:
    text = re.sub(r"<style.*?</style>", "", html_content, flags=re.I | re.S)
    text = re.sub(r"<script.*?</script>", "", text, flags=re.I | re.S)
    text = re.sub(r"</(p|h1|h2|h3|li|div|tr)>", "\n", text, flags=re.I)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return unescape(re.sub(r"\n{3,}", "\n\n", text)).strip()


def report_pdf_bytes(
    title: str,
    content_html: str,
    charts: Iterable[Any],
    metadata: dict[str, Any],
) -> bytes:
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title,
    )
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(title, styles["Title"]), Spacer(1, 0.4 * cm)]
    story.append(Paragraph("Rascunho para revisão científica e pedagógica.", styles["Italic"]))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(f"Dataset: {metadata.get('dataset_checksum', '—')}", styles["BodyText"]))
    story.append(Paragraph(f"Análise: {metadata.get('analysis_signature', '—')}", styles["BodyText"]))
    story.append(Spacer(1, 0.5 * cm))
    for paragraph in _plain_text(content_html).split("\n"):
        if paragraph.strip():
            story.append(Paragraph(escape(paragraph.strip()), styles["BodyText"]))
            story.append(Spacer(1, 0.15 * cm))
    for chart in charts:
        image_bytes = render_chart_bytes(
            chart.chart_type,
            chart.data_snapshot,
            chart.configuration,
            chart.title,
            chart.description,
            "png",
        )
        story.extend(
            [
                PageBreak(),
                Paragraph(chart.title, styles["Heading2"]),
                Paragraph(chart.alt_text, styles["BodyText"]),
                Spacer(1, 0.2 * cm),
                Image(io.BytesIO(image_bytes), width=17 * cm, height=9.8 * cm),
            ]
        )
    document.build(story)
    return output.getvalue()


def report_docx_bytes(
    title: str,
    content_html: str,
    charts: Iterable[Any],
    metadata: dict[str, Any],
) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph("Rascunho para revisão científica e pedagógica.")
    document.add_paragraph(f"Dataset SHA-256: {metadata.get('dataset_checksum', '—')}")
    document.add_paragraph(f"Assinatura da análise: {metadata.get('analysis_signature', '—')}")
    for paragraph in _plain_text(content_html).split("\n"):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
    for chart in charts:
        document.add_heading(chart.title, level=2)
        document.add_paragraph(chart.alt_text)
        image_bytes = render_chart_bytes(
            chart.chart_type,
            chart.data_snapshot,
            chart.configuration,
            chart.title,
            chart.description,
            "png",
        )
        document.add_picture(io.BytesIO(image_bytes), width=Inches(6.4))
    document.add_heading("Metadados de reprodutibilidade", level=2)
    document.add_paragraph(json.dumps(metadata, ensure_ascii=False, indent=2))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
